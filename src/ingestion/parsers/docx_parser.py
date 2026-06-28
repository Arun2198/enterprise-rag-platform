from pathlib import Path

from docx import Document as DocxDocument

from ingestion.contracts.document import Document
from ingestion.contracts.result import Error
from ingestion.contracts.result import Result
from ingestion.parsers.base_parser import BaseParser


class DOCXParser(BaseParser):

    def parse(
        self,
        file_path: str
    ) -> Result[Document]:

        try:
            path = Path(file_path)

            if not path.exists():
                return Result(
                    success=False,
                    error=Error(
                        code="FILE_NOT_FOUND",
                        message=f"{file_path} does not exist"
                    )
                )

            docx = DocxDocument(path)
            text_blocks = []

            for paragraph in docx.paragraphs:
                text = paragraph.text.strip()
                if text:
                    text_blocks.append(text)

            for table in docx.tables:
                for row in table.rows:
                    row_text = [
                        cell.text.strip()
                        for cell in row.cells
                        if cell.text.strip()
                    ]
                    if row_text:
                        text_blocks.append(" | ".join(row_text))

            content = "\n".join(text_blocks)

            if not content.strip():
                return Result(
                    success=False,
                    error=Error(
                        code="EMPTY_DOCUMENT",
                        message="No text extracted from DOCX"
                    )
                )

            document = Document(
                document_id=path.stem,
                source=str(path),
                document_type="docx",
                content=content,
                metadata={
                    "paragraph_count": len(docx.paragraphs),
                    "table_count": len(docx.tables),
                    "file_name": path.name
                }
            )

            return Result(
                success=True,
                data=document
            )

        except Exception as ex:
            return Result(
                success=False,
                error=Error(
                    code="DOCX_PARSE_ERROR",
                    message=str(ex)
                )
            )
