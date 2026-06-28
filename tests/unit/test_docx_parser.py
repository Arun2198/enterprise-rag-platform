from docx import Document

from ingestion.parsers.docx_parser import DOCXParser


def test_docx_parser_success(tmp_path):

    file_path = tmp_path / "policy.docx"
    docx = Document()
    docx.add_paragraph("Leave Policy 2025")
    table = docx.add_table(rows=1, cols=2)
    table.cell(0, 0).text = "Employees"
    table.cell(0, 1).text = "20 days"
    docx.save(file_path)

    parser = DOCXParser()

    result = parser.parse(str(file_path))

    assert result.success is True
    assert result.data is not None
    assert "Leave Policy 2025" in result.data.content
    assert "Employees | 20 days" in result.data.content
    assert result.data.metadata["table_count"] == 1


def test_docx_parser_missing_file():

    parser = DOCXParser()

    result = parser.parse("sample_documents/does_not_exist.docx")

    assert result.success is False
    assert result.error is not None
    assert result.error.code == "FILE_NOT_FOUND"
